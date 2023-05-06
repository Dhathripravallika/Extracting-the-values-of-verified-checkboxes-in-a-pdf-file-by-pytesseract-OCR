import numpy as np
import pytesseract
import pandas as pd
import cv2
from pdf2image import convert_from_path
import os
from PIL import Image
from pytesseract import image_to_string
from module import extract_boxes, extract_checkboxes, read_docx_table
from docx import Document
from datetime import datetime
start_time = datetime.now()



images = convert_from_path('C:/Users/Taraneh/Desktop/pythonProject/ocr/pages/typeA.pdf',
                           poppler_path='C:/poppler-0.68.0/bin')
path = "pages"

for i in range(len(images)):
    name = f'page' + str(i) + '.jpg'
    filename = os.path.join(path, name)  # Save pages as images in the pdf
    images[i].save(filename)



wordfile = input("Enter path of input file with .docx suffix: ")
document = Document(wordfile)

df_main = read_docx_table(document,table_num=1,nheader=1)

desired_param = {'indx':[0,1,2,3,4,5,6,7,], 'values':['Temperature', 'humidity ', 'Altitude', 'Pollution', 'protection', 'Mass', 'Dimensions' , 'Overvoltage',]}
desired_param_df = pd.DataFrame(desired_param)

output = []
text_list2 = []
output2 = []
output3 = []

for index, parameter in desired_param_df.values:
    output = []
    output2 = []
    text_list2 = []
    for root, dirs, files in os.walk(path):
        for k, filename in enumerate(files):
            pg_nmb = filename.split(".")

            if pg_nmb[1] == 'jpg':
                pg_nmb = pg_nmb[0]
                file_path = os.path.join(root, filename)
                results = pytesseract.image_to_data(file_path, output_type=pytesseract.Output.DICT,
                                                    config='--psm 12 --oem 3 ')
                results_df = pd.DataFrame(results)
                checkboxes = extract_checkboxes(file_path)
                boxes = extract_boxes(file_path)

                row_of_parameter = np.where(results_df["text"] == parameter)
                # if row_of_parameter[0].size>0:
                #   rows_of_parameters.append(row_of_parameter[0][0])  #in the working file
                # for num, row in enumerate(rows_of_parameters):
                parmtrs_block = results_df['block_num'].iloc[row_of_parameter]  # block of the parameter in the working file
                values_block = [parmtrs_block + 1]  # block of the value corresonding to the block

                values_rows = np.where(
                    results_df['block_num'].isin(values_block[0]))  # row of the values in the working file
                output_text = " ".join(results_df['text'].iloc[values_rows])
                if output_text != '':
                    output.append(output_text)
                df_main['output'].iloc[index] = output

                if values_rows[0].size > 0:
                    xmin, ymin = results_df['left'].iloc[values_rows[0][0]], results_df['top'].iloc[
                        values_rows[0][0]]  # coordinates of fist token in corresponding block
                    xmax = results_df['left'].iloc[values_rows[0][-1]] + results_df['width'].iloc[values_rows[0][-1]]
                    ymax = results_df['top'].iloc[values_rows[0][-1]] + results_df['height'].iloc[
                        values_rows[0][-1]]  # coordinates of the last token

                    for ff, box in enumerate(boxes):

                        if (box[0] < xmin < box[0] + box[2]) and (box[1] < ymin < box[1] + box[3]):
                            # and boxes[ff+1][1]==box[1]:#block is in the box(box after parameter) and it has box ghablesh
                            output_box = box
                            for k, checkbox in enumerate(checkboxes):
                                if output_box[0] < checkbox[0][0] < output_box[0] + output_box[2] and output_box[1] < \
                                        checkbox[0][1] < output_box[1] + output_box[3] and checkbox[1] == 1:  # if there is checkbox in the corresponding block(next block after parameters)

                                    for j in range(0, len(results_df)):
                                        x, y, w, h = results_df["left"][j], results_df["top"][j], results_df["width"][j], \
                                        results_df["height"][j]

                                        if (k + 1) != len(checkboxes):
                                            next_box = checkboxes[k + 1]
                                            if checkbox[0][1] - 10 <= next_box[0][1] <= checkbox[0][1] + 10:
                                                if (checkbox[0][1] <= y <= checkbox[0][1] + checkbox[0][2]) and (
                                                        checkbox[0][0] + checkbox[0][2] <= x <= next_box[0][0]):
                                                    text_list2.append(results_df['text'][j])
                                            else:
                                                if (checkbox[0][1] <= y <= checkbox[0][1] + checkbox[0][2]) and (
                                                        checkbox[0][0] + checkbox[0][2] <= x):
                                                    text_list2.append(results_df['text'][j])

                                        if (k + 1) == len(checkboxes) and (
                                                checkbox[0][1] <= y <= checkbox[0][1] + checkbox[0][2]) and (
                                                checkbox[0][0] + checkbox[0][2] <= x):
                                            text_list2.append((results_df['text'][j]))

                                    output2.append(text_list2)



                                elif not any(
                                        output_box[0] < checkbox[0][0] < output_box[0] + output_box[2] and output_box[1] <
                                        checkbox[0][1] < output_box[1] + output_box[3] for checkbox in checkboxes):
                                    output2.append(" ".join(results_df['text'].iloc[values_rows[0]]))
                                df_main['output'].iloc[index] = output2

            # output2.append(text_list2)
            # output2.append(output)
            # df_main['output'].iloc[index] = output
            # print(df_main)
        #   df_main['output'].iloc[index] = output2
        #   print(df_main)
print(df_main)
    #     # values_block = [parmtr_block+1 for parmtr_block in parmtrs_block]


# open an existing document
doc = Document(wordfile)
# add a table to the end and create a reference variable
# extra row is so we can add the header row
t = doc.add_table(df_main.shape[0]+1, df_main.shape[1])
# add the header rows.
for j in range(df_main.shape[-1]):
    t.cell(0,j).text = df_main.columns[j]
# add the rest of the data frame
for i in range(df_main.shape[0]):
    for j in range(df_main.shape[-1]):
        t.cell(i+1,j).text = str(df_main.values[i,j])
# save the doc
doc.save(wordfile)


end_time = datetime.now()
print('Duration: {}'.format(end_time - start_time))